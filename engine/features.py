import os
from pipes import quote
import re
import sqlite3
import struct
import subprocess
import time
import webbrowser
from playsound import playsound
import eel
import pyaudio
import pyautogui
from engine.command import speak
from engine.config import ASSISTANT_NAME
import pywhatkit as kit
import pvporcupine
from hugchat import hugchat
import openpyxl
from openpyxl import load_workbook

from engine.helper import extract_yt_term, remove_words
import openai
import docx
import platform
import pyttsx3


con = sqlite3.connect("vasu.db")
cursor = con.cursor()

#playing assistant sound function
@eel.expose
def playAssistantsound():
    music_dir = "www\\assets\\audio\\start_sound.mp3"
    playsound(music_dir)



def openCommand(query):
    query = query.replace(ASSISTANT_NAME, "")
    query = query.replace("open", "")
    query.lower()

    app_name = query.strip()

    if app_name != "":

        try:
            cursor.execute(
                'SELECT path FROM sys_command WHERE name IN (?)', (app_name,))
            results = cursor.fetchall()

            if len(results) != 0:
                speak("Opening "+query)
                os.startfile(results[0][0])
                

            elif len(results) == 0: 
                cursor.execute(
                'SELECT url FROM web_command WHERE name IN (?)', (app_name,))
                results = cursor.fetchall()
                
                if len(results) != 0:
                    speak("Opening "+query)
                    webbrowser.open(results[0][0])

                else:
                    speak("Opening "+query)
                    try:
                        os.system('start '+query)
                    except:
                        speak("not found")
        except:
            speak("some thing went wrong")





        
  

# Function to save content to Word document
def save_to_word(content, filename):
    downloads_dir = os.path.join(os.path.expanduser('~'), 'Downloads')
    filepath = os.path.join(downloads_dir, filename)
    if os.path.exists(filepath):
        doc = docx.Document(filepath)
    else:
        doc = docx.Document()
    doc.add_paragraph(content)
    doc.save(filepath)

# Function to create a new Word file in Downloads directory
def create_new_word_file(filename):
    downloads_dir = os.path.join(os.path.expanduser('~'), 'Downloads')
    filepath = os.path.join(downloads_dir, filename)
    # Create an empty Word document
    doc = docx.Document()
    doc.save(filepath)
    print(f"New Word file created: {filepath}")
    return filepath



# Function to read content from a Word file in Downloads directory
def read_word_file(filename):
    downloads_dir = os.path.join(os.path.expanduser('~'), 'Downloads')
    filepath = os.path.join(downloads_dir, filename)
    if os.path.exists(filepath):
        try:
            os.startfile(filepath)  # Open the Word file
        except OSError:
            print("Unable to open the Word file.")
            return None

        doc = docx.Document(filepath)
        text = ''
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
        return text
    else:
        print("File not found.")
        return None

def speak_text(text):
    engine = pyttsx3.init()
    engine.say(text)
    engine.runAndWait()


def PlayYoutube(query):
    search_term = extract_yt_term(query)
    speak("Playing "+search_term+" on YouTube")
    kit.playonyt(search_term)


# find contacts
def findContact(query):
    
    words_to_remove = [ASSISTANT_NAME, 'make', 'a', 'to', 'phone', 'call', 'send', 'message', 'whatsapp', 'video']
    query = remove_words(query, words_to_remove)

    try:
        query = query.strip().lower()
        cursor.execute("SELECT mobile_no FROM contacts WHERE LOWER(name) LIKE ? OR LOWER(name) LIKE ?", ('%' + query + '%', query + '%'))
        results = cursor.fetchall()
        print(results[0][0])
        mobile_number_str = str(results[0][0])

        if not mobile_number_str.startswith('+91'):
            mobile_number_str = '+91' + mobile_number_str

        return mobile_number_str, query
    except:
        speak('not exist in contacts')
        return 0, 0
    
def whatsApp(mobile_no, message, flag, name):
    

    if flag == 'message':
        target_tab = 12
        jarvis_message = "message send successfully to "+name

    elif flag == 'call':
        target_tab = 7
        message = ''
        jarvis_message = "calling to "+name

    else:
        target_tab = 6
        message = ''
        jarvis_message = "staring video call with "+name


    # Encode the message for URL
    encoded_message = quote(message)
    print(encoded_message)
    # Construct the URL
    whatsapp_url = f"whatsapp://send?phone={mobile_no}&text={encoded_message}"

    # Construct the full command
    full_command = f'start "" "{whatsapp_url}"'

    # Open WhatsApp with the constructed URL using cmd.exe
    subprocess.run(full_command, shell=True)
    time.sleep(5)
    subprocess.run(full_command, shell=True)
    
    pyautogui.hotkey('ctrl', 'f')

    for i in range(1, target_tab):
        pyautogui.hotkey('tab')

    pyautogui.hotkey('enter')
    speak(jarvis_message)  


    # chat bot 
def chatBot(query):
    user_input = query.lower()
    chatbot = hugchat.ChatBot(cookie_path="engine\cookies.json")
    id = chatbot.new_conversation()
    chatbot.change_conversation(id)
    response =  chatbot.chat(user_input)
    print(response)
    speak(response)
    return response  


def open_excel_file(filename):
    # Remove the word "file" from the filename
    filename = filename.replace("file", "").strip()
    
    downloads_folder = os.path.expanduser("C:/Users/SUJAL/Downloads")  # Change this to your Downloads folder path
    for file in os.listdir(downloads_folder):
        if file.startswith(filename) and file.endswith(".xlsx"):
            file_path = os.path.join(downloads_folder, file)
            os.startfile(file_path)
            print("Excel file '{}' opened successfully.".format(file))
            return True
    print("File '{}' not found in the Downloads folder.".format(filename))
    return False

    
def chart_creation_file(command):
    filename = command.split("data in")[-1].strip()
    found_file = False
    file_found = None
    downloads_folder = os.path.expanduser("C:/Users/SUJAL/Downloads")
    for file in os.listdir(downloads_folder):
        if file.startswith(filename) and file.endswith(".xlsx"):
            file_found = os.path.join(downloads_folder, file)
            found_file = True
            break
    
    if found_file:
        wb = load_workbook(file_found)
        sheet = wb.active
        if has_data(sheet):
            create_charts(file_found, sheet)
        else:
            speak("No data found in the Excel file. Please add data to create charts.")
    else:
        speak("File '{}' not found in the Downloads folder.".format(filename))

def create_new_excel_file(filename, content):
    default_save_location = os.path.expanduser("C:/Users/SUJAL/Downloads")
    filepath = os.path.join(default_save_location, filename + '.xlsx')
    wb = openpyxl.Workbook()
    sheet = wb.active
    sheet['A1'] = content
    wb.save(filepath)
    speak("Excel file '{}' created successfully and saved at: {}".format(filename, filepath))

def has_data(sheet):
    for row in sheet.iter_rows():
        for cell in row:
            if cell.value:
                return True
    return False

def create_charts(filepath, sheet):
    os.startfile(filepath)
    time.sleep(3)
    pyautogui.hotkey('winleft', 'up')
    if not has_data(sheet):
        speak("No data found in the Excel file. Please add data to create charts.")
        return

    rows = sheet.max_row
    cols = sheet.max_column
    start_cell = None
    end_cell = None
    for row in range(1, rows + 1):
        for col in range(1, cols + 1):
            if sheet.cell(row=row, column=col).value:
                if start_cell is None:
                    start_cell = (col, row)
                end_cell = (col, row)

    if start_cell and end_cell:
        start_col, start_row = start_cell
        end_col, end_row = end_cell
        start_cell_address = openpyxl.utils.get_column_letter(start_col) + str(start_row)
        end_cell_address = openpyxl.utils.get_column_letter(end_col) + str(end_row)
        range_address = start_cell_address + ":" + end_cell_address
        pyautogui.hotkey('ctrl', 'g')
        pyautogui.write(range_address)
        pyautogui.press('enter')
        pyautogui.hotkey('alt', 'n')
        pyautogui.moveTo(605, 144, duration=0.5)
        pyautogui.click()
        time.sleep(1)
        pyautogui.moveTo(713, 377, duration=0.5)
        pyautogui.press('enter')
        time.sleep(1)
        speak("Bar chart created successfully.")
    else:
        speak("No data found to create a chart.")




def navigate(direction):
    if direction == "up":
        pyautogui.press('up')
    elif direction == "down":
        pyautogui.press('down')
    elif direction == "left":
        pyautogui.press('left')
    elif direction == "right":
        pyautogui.press('right')
    else:
        speak("Invalid direction. Please try again.")

def go_to_cell(cell_reference):
    cell_reference = cell_reference.upper()
    column, row = '', ''
    for char in cell_reference:
        if char.isalpha():
            column += char
        elif char.isdigit():
            row += char
        else:
            speak("Invalid cell reference format. Please provide a valid cell reference.")
            return

    row_number = int(row) - 1
    column_index = ord(column) - ord('A')

    pyautogui.click(x=100, y=100)
    pyautogui.hotkey('f5')
    pyautogui.write(column + row)
    pyautogui.press('enter')
    time.sleep(1)

def write_data(data):
    pyautogui.write(data)

def delete_data():
    pyautogui.hotkey('delete')